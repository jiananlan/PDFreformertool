from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Pt
import io
from docx.oxml.ns import qn
from PIL import Image
from docx.oxml import OxmlElement
import xml.etree.ElementTree as ET
import rich
import T5_a
import re


def get_rid(txt):
    pattern = r'r:embed="(.*?)"/>'
    matches = re.findall(pattern, txt)
    return matches


def get_size(text):
    pattern = r'<wp:extent cx="([^"]+)" cy="([^"]+)"/>'
    matches = re.findall(pattern, text)
    matches = [int(m) for m in list(matches[0])]
    print(matches)
    return matches


def process_shell(file, to):
    def get_cell_font_size(cell):
        """尝试获取单元格内文本的字体大小"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.font.size:
                    print("🐞I have got the size")
                    return run.font.size  # 返回第一个 `Run` 里定义的字体大小
        print("🥏 No size")
        return Pt(7)  # 默认返回 12 磅

    def judege(text):
        """

        :param text:输入的文字
        :return: 0:无需翻译
                    1:需要简单翻译
                    2:需要正常翻译
        """
        try:
            eval(text)
            return 0
        except:
            if len(text.replace(" ", "")) > 40:
                return 2
            else:
                return 1

    def process_text(text):
        return T5_a.query_key(text)

    # 读取 Word 文件
    doc = Document(file)

    list_of_simple_trs = set()
    list_of_full_trs = set()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    if judege(cell.text) == 1:
                        list_of_simple_trs.add(cell.text)
                    elif judege(cell.text) == 2:
                        list_of_full_trs.add(cell.text)
    print(list_of_simple_trs)
    print(list_of_full_trs)
    list_of_simple_trs = list(list_of_simple_trs)
    list_of_full_trs = list(list_of_full_trs)
    import T4_b
    import asyncio

    def set_vertical_center(cell):
        """设置单元格内容垂直居中"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # 创建 <w:vAlign> 元素，并正确设置命名空间
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")  # 确保使用正确的 XML 命名空间

        tcPr.append(vAlign)

    def set_left_and_right_indent(cell, left_indent_pt=5, right_indent_pt=5):
        """设置单元格文本左对齐并且左右缩进"""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 文字左对齐
            paragraph.paragraph_format.left_indent = Pt(left_indent_pt)  # 设置左侧缩进
            paragraph.paragraph_format.right_indent = Pt(
                right_indent_pt
            )  # 设置右侧缩进

    asyncio.run(T4_b.run((list_of_simple_trs, list_of_full_trs)))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    if judege(cell.text):
                        new_text = process_text(cell.text)
                    else:
                        new_text = cell.text
                    fts = get_cell_font_size(cell)
                    flag = False
                    image_rid = ""
                    size_data = ""
                    for index_1, para in enumerate(cell.paragraphs):
                        for index_2, run in enumerate(para.runs):
                            if "a:graphic" in run._r.xml:
                                print(f"[{index_1, index_2}]🎈", run._r.xml)
                                flag = True
                                image_rid = get_rid(run._r.xml)[0]
                                size_data = get_size(run._r.xml)
                    for p in cell.paragraphs:
                        p.text = ""  # if p.text else p.text
                    if flag:
                        with open("data.txt", "r") as f:
                            t = eval(f.read())
                        with open("size.txt", "w") as f:
                            f.write(f'"{t[image_rid]}",{size_data[0]},{size_data[1]}\n')
                        print("🍎", image_rid, size_data, t[image_rid])
                        if cell.paragraphs and cell.paragraphs[0].runs:

                            print(f"Add {t[image_rid].split('.')[0]}")
                            cell.paragraphs[0].add_run(
                                "{{" + t[image_rid].split(".")[0] + "}}"
                            )
                    paragraph = cell.paragraphs[0]  # 获取单元格中的第一个段落
                    run = paragraph.add_run(new_text)  # 添加文本
                    run.font.size = fts
                    set_vertical_center(cell)
                    set_left_and_right_indent(cell, 10)

    doc.save(to)
