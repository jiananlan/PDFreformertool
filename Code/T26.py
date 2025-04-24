from docx import Document
import T5_a
from docx.oxml.ns import qn
import os
import asyncio
import T4_a
import T24
import Tconfig


def process_text(text):
    """
    这是一个示例函数，用于处理文本内容。
    你可以根据需要替换为你的实际处理逻辑，例如翻译。
    """
    print(f'->{text}')
    return T5_a.query_key(key=text)


def modify_paragraphs(docx_path, sec_docx_path, process_paragraph_function):
    """
    遍历DOCX文件中的所有段落，将段落中的所有文字合并后使用
    process_paragraph_function处理，然后将处理后的文本写回段落，
    并尽可能保留原有格式。非文字内容（如图片）的段落保持不变。

    Args:
        docx_path (str): DOCX文件的路径。
        process_paragraph_function (callable): 用于处理段落文本的函数。
    """

    def mode(data):
        """计算列表data中的众数，并返回包含所有众数的列表"""
        if not data:
            return []  # 列表为空时返回空列表

        counts = {}
        for num in data:
            counts[num] = counts.get(num, 0) + 1  # 统计每个数字出现的次数

        max_count = max(counts.values())  # 找出最大的出现次数

        modes = [k for k, v in counts.items() if v == max_count]  # 收集所有众数

        return modes

    document = Document(docx_path)

    document.styles["Normal"].font.name = "Times New Roman"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    set_of_raw_content = set()
    list_of_font_size = []
    for paragraph in document.paragraphs:
        text_runs = [run for run in paragraph.runs if run.element.tag.endswith('}r') and run.element.xpath('./w:t')]
        if text_runs:
            original_paragraph_text = "".join(run.text for run in text_runs)
            set_of_raw_content.add(original_paragraph_text)
            list_of_font_size.append(text_runs[0].font.size)
    mode_size = mode(list_of_font_size)
    set_of_raw_content = list(set_of_raw_content)
    list_of_raw_content = list()
    for raw in set_of_raw_content:
        if not T5_a.key_exs(raw) and len(raw) > 10:
            list_of_raw_content.append(raw)
    print(list_of_raw_content)

    if not Tconfig.configuration["enable_chatgpt"]:
        asyncio.run(T4_a.run(prompts=[x for x in list_of_raw_content if len(x) > 10]))
    else:
        asyncio.run(T24.main([x for x in list_of_raw_content if len(x) > 10]))

    for paragraph in document.paragraphs:
        text_runs = [run for run in paragraph.runs if run.element.tag.endswith('}r') and run.element.xpath('./w:t')]
        if text_runs:
            original_paragraph_text = "".join(run.text for run in text_runs)
            processed_paragraph_text = process_paragraph_function(original_paragraph_text)
            if paragraph.runs:

                first_run_format = paragraph.runs[0].font
                paragraph.clear()
                new_run = paragraph.add_run(processed_paragraph_text)
                if first_run_format.size:
                    print("🐞I have got the size")
                    if first_run_format.size>max(mode_size):

                        print("😹 bigger")
                        new_run.font.name = "Times New Roman"
                        new_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                    elif first_run_format.size<min(mode_size):
                        print("🙈 smaller")
                        new_run.font.name = "Times New Roman"
                        new_run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                    else:
                        new_run.font.name = first_run_format.name
                else:
                    new_run.font.name = first_run_format.name
                new_run.font.size = first_run_format.size
                new_run.font.bold = first_run_format.bold
                new_run.font.italic = first_run_format.italic
                new_run.font.underline = first_run_format.underline
                new_run.font.color.rgb = first_run_format.color.rgb
                # Add more formatting attributes as needed

    document.save(sec_docx_path)
    print(f"已按段落处理并保存为 '{sec_docx_path}'")


if __name__ == '__main__':
    file_path = 'H-your_document.docx'  # 替换为你的实际文件路径


    def translate_example(text):
        """
        一个简单的翻译示例 (仅用于演示).
        在实际应用中，你需要使用真正的翻译库或服务。
        """
        translations = {
            "Hello": "你好",
            "world": "世界",
            "This is a test": "这是一个测试"
        }
        return translations.get(text, text)


    # 使用示例的翻译函数
    modify_paragraphs('hT_iFWpPzfCV.docx', file_path, process_text)
    os.startfile(filepath=file_path)
