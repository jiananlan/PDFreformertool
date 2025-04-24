from docx import Document
from docx.oxml.ns import qn
import rich
import T5_a
import asyncio
import T10
import T4_a
import os
import T24
import Tconfig

count = 0


def main_run(doc_path, third):
    global count
    # Load the docx file
    doc = Document(doc_path)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Function to manipulate text (you can replace this with any function)

    def manipulate_text(text):
        global count
        if len(text.replace(" ", "")) < 4 or text.replace(".", "").isnumeric():
            return text
        count += 1
        return T5_a.query_key(key=text)

    list_of_raw_content = set()
    for IDX, para in enumerate(doc.paragraphs):
        print(f"[{IDX + 1}]->", para.text)
        list_of_raw_content.add(para.text)
    list_of_raw_content = list(list_of_raw_content)
    list_of_raw_content.sort(key=lambda x: len(x))
    # list_of_raw_content = list_of_raw_content[::-1]
    if not Tconfig.configuration["enable_chatgpt"]:
        asyncio.run(T4_a.run(prompts=[x for x in list_of_raw_content if len(x) > 10]))
    else:
        asyncio.run(T24.main([x for x in list_of_raw_content if len(x) > 10]))

    current = 1
    total = 0

    def get_image(num):
        global current, total
        total += num
        highest = num + current
        img_list = os.listdir(".\\extracted_images")
        result = []
        for filename in img_list:
            if filename.startswith(f"image{current}"):
                result.append(os.path.join(".\\extracted_images", filename))
                current += 1
                if current >= highest:
                    return result

    def mode(data):
        """计算列表data中的众数，并返回包含所有众数的列表"""
        if not data:
            return []  # 列表为空时返回空列表

        counts = {}
        for num in data:
            counts[num] = counts.get(num, 0) + 1  # 统计每个数字出现的次数

        max_count = max(counts.values())  # 找出最大的出现次数

        modes = [k for k, v in counts.items() if v == max_count]  # 收集所有众数

        return modes

    mode_size = []

    for para in doc.paragraphs:
        if para.runs and para.runs[0].font.size:
            mode_size.append(float(para.runs[0].font.size))

    mode_size = mode(mode_size)
    print(mode_size)
    t, i = T10.read_word_docx(doc_path)
    print(len(i), i)
    i1 = [cc[0] for cc in i]
    for idx, para in enumerate(doc.paragraphs):
        for x in i:
            if x[0] == idx + 1:
                print(x[-1])
                para.runs[0].add_text("{{" + x[-1].split(".")[0] + "}}")
            """    if idx + 1 in i1:
            pic_add = i1.count(idx + 1)
            print(idx + 1, i1.count(idx + 1), pic_add)
            for pic in pic_add:
                para.runs[0].add_text('{{' + pic.split('\\')[-1].split('.')[0] + '}}')
            # 添加图片"""
        if para.runs:
            first_run = para.runs[0]
            new_text = manipulate_text(para.text)
            para.clear()
            run = para.add_run(new_text)
            run.bold = first_run.bold
            run.italic = first_run.italic
            run.underline = first_run.underline
            run.font.size = first_run.font.size
            run.font.color.rgb = first_run.font.color.rgb
            if first_run.font.size:
                if first_run.font.size > max(mode_size):
                    rich.print("[🥝 Notice:]This is bigger")
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                elif first_run.font.size < min(mode_size):
                    rich.print("[🥥 Notice:]This is smaller")
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    print(count)
    print(len(i), total, i)
    doc.save(third)

    print(os.listdir(".\\extracted_images"))
    print(mode_size)
    """
    
    """
    return T4_a.error_count


if __name__ == "__main__":
    main_run("hT_iFWpPzfCV.docx", 'nothing.docx')
