from pdf2docx import parse
from docx import Document

pdf_file = r"D:\电脑管家迁移文件\微信聊天记录搬家\WeChat Files\wxid_wbjxn8hzlq6w22\FileStorage\File\2024-11\科研入门培养方案\科研入门培养方案\MULTI_Manual.pdf"  # 替换为你的PDF文件路径
docx_file = "output3.docx"


def run(pdf_file, docx_file):
    parse(pdf_file, docx_file, start=0, end=None)
if __name__ == "__main__":
    run(pdf_file, docx_file)

"""doc = Document(docx_file)
doc.paragraphs[0].text = "🚀 Edited First Paragraph: " + doc.paragraphs[0].text  # Modify first paragraph

# Add a new paragraph
doc.add_paragraph("This is a new paragraph added after conversion.")

# Step 4: Save the modified document
doc.save("edited_output.docx")

print("DOCX file edited and saved as 'edited_output.docx'")"""
